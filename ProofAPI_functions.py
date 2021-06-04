import sys
import os
import re
from docx import Document
from PIL import Image
import pandas as pd


def main(input_directory):
    article = input_directory.split("\\")[-1]
    data = {'Manuscript ID': [],
            'Title (word count)': [],
            'Abstract (word count)': [],
            'KeyWords (word count)': [],
            'Main Titles': [],
            'Article (word count)': [],
            'Tables': [],
            'Figures': [],
            'Figures resolutions': [],
            'Direct quotations': [],
            'In-text references': [],
            'Reference list': []}
    df = pd.DataFrame(data)
    document = Document(os.path.join(input_directory, '{0}.docx'.format(article)))
    print('Checking {0}'.format(article))
    # Check front (word counts of manuscript title, abstract and keywords, author and affiliation information)
    try:
        front_report = check_front(document)
    except Exception as e:
        print('Error in checking front: {0}'.format(str(e)))
    # Check body (main titles)
    try:
        body_report = check_body(document)
    except Exception as e:
        print('Error in checking body: {0}'.format(str(e)))
    # Check tables order
    try:
        tab_fig_report = check_tables_and_figures(document)
        tables = ''
        figures = ''
        for key in tab_fig_report:
            if 'Table' in key:
                tables += key + ': ' + tab_fig_report[key] + ', '
            else:
                figures += key + ': ' + tab_fig_report[key] + ', '
        tables = tables.rstrip(', ')
        figures = figures.rstrip(', ')
    except Exception as e:
        print('Error in checking tables & figures: {0}'.format(str(e)))
    # Check figures quality
    try:
        fig_quality = check_figures_resolutions(input_directory)
        figures_quality = ''
        for key in fig_quality:
            figures_quality += key + ': ' + fig_quality[key] + ', '
        figures_quality = figures_quality.rstrip(', ')
    except Exception as e:
        print('Error in checking figures quality: {0}'.format(str(e)))
    # Check direct quotations
    try:
        quotations_report = check_direct_quotations(document)
        quotations = ''
        for key in quotations_report:
            quotations += key + ': ' + quotations_report[key] + ', '
        quotations = quotations.rstrip(', ')
    except Exception as e:
        print('Error in checking quotations: {0}'.format(str(e)))
    # Check references
    try:
        ref_report = check_references(document)
        references = ''
        for key in ref_report:
            if key != 'Order':
                references += key + ': ' + ref_report[key] + ', '
        references = references.rstrip(', ')
    except Exception as e:
        print('Error in checking references: {0}'.format(str(e)))
    try:
        df = df.append({'Manuscript ID': article, 'Title (word count)': front_report['Title'],
                        'Abstract (word count)': front_report['Abstract']
                           , 'KeyWords (word count)': front_report['KeyWords'],
                        'Main Titles': body_report['mainHeaderOrder'],
                        'Article (word count)': body_report['wordCount']
                           , 'Tables': tables, 'Figures': figures, 'Figures resolutions': figures_quality,
                        'Direct quotations': quotations
                           , 'In-text references': references, 'Reference list': ref_report['Order']},
                       ignore_index=True)
        # generate_report(front_report, body_report, tab_fig_report, fig_quality, quotations_errors, ref_report)
    except Exception as e:
        print('Error in generating the report: {0}'.format(str(e)))
    print('Done')
    return df.to_json(orient='records')


def check_front(document):
    front = []
    front_report = {}
    for para in document.paragraphs:
        if para.style.style_id == 'FA-MainHeader':
            break
        else:
            front.append(para)
    # TITLE (Not more than 12 words)
    title = []
    [title.append(item) for item in front if (item.style.style_id == 'FA-PaperTitle')]
    if len(title) > 1:
        print('There is more than one title, it must be merged into one paragraph !!')
        sys.exit()
    front_report['Title'] = len(title[0].text.split())
    # ABSTRACT (Not more than 250 words) (minimum = 150)
    abstract = []
    [abstract.append(item) for item in front if (item.style.style_id == 'FA-Paragraphtext')]
    if len(abstract) > 1:
        print('There is more than one paragraph in abstract, it must be merged into one paragraph !!')
        sys.exit()
    front_report['Abstract'] = len(abstract[0].text.split())
    # KEYWORDS (Not more than 6 keywords) (add check for the content of the keywords)
    key_words = []
    [key_words.append(item) for item in front if (item.style.style_id == 'FA-RefHeader')]
    front_report['KeyWords'] = len(key_words[0].text.split(','))
    return front_report


def check_body(document):
    # ARTICLE_WORD_COUNT (Not more than 5000)
    words = 0
    for paragraph in document.paragraphs:
        words = words + len(paragraph.text.split())
    # ("introduction", "problem statement", "research questions", "purpose of the study", "research methods", "findings", "conclusion", "references")
    main_header_order = ["introduction", "problem statement", "research questions", "purpose of the study",
                         "research methods", "findings", "conclusion"]
    main_headers = [f.text.lower() for f in document.paragraphs if f.style.name == "FA-Main Header"]
    body_report = {
        "wordCount": words,
        "mainHeaderOrder": "ok" if (main_header_order == main_headers) else "Error"
    }
    return body_report


def check_tables_and_figures(document):
    tab_fig_report = {}
    tables = {}
    figures = {}
    tables_ref = []
    figures_ref = []
    tindex = 0
    findex = 0
    index = 0
    checked_tab = []
    checked_fig = []
    for item in document.paragraphs:
        if item.style.style_id == 'FA-FNumber':
            tables[tindex + 1] = item
            tindex += 1
        elif item.style.style_id == 'FA-Fnumber0':
            figures[findex + 1] = item
            findex += 1
        elif item.style.style_id == 'FA-Paragraphtext':
            curent_ref = []
            content = item.text
            for match in re.findall(r'(Table\s*\d+(\-\d+)*|Tab\s*\d+(\-\d+)*)', content):
                if ',' in match[0]:
                    [curent_ref.append(int(x.replace('Table ', '').replace('Tab ', ''))) for x in match[0].split(',')]
                    [tables_ref.append(int(x.replace('Table ', '').replace('Tab ', ''))) for x in match[0].split(',')]
                    [curent_ref.remove(tab) for tab in curent_ref if tab in checked_tab]
                else:
                    curent_ref.append(int(match[0].replace('Table ', '').replace('Tab ', '')))
                    tables_ref.append(int(match[0].replace('Table ', '').replace('Tab ', '')))
                    [curent_ref.remove(tab) for tab in curent_ref if tab in checked_tab]
            curent_ref = list(set(curent_ref))
            refers_count = len(curent_ref)
            if refers_count > 0:
                counter = 0
                for i in range(index + 1, len(document.paragraphs)):
                    if document.paragraphs[i].style.style_id != 'Normal' and document.paragraphs[
                        i].style.style_id != 'FA-FNumber':
                        tab_fig_report['Table_{0}'.format(curent_ref[counter])] = 'Refer not places correctly'
                    else:
                        if document.paragraphs[i].style.style_id == 'FA-FNumber':
                            counter += 1
                    if counter == refers_count:
                        break
                [checked_tab.append(x) for x in curent_ref]
            curent_ref = []
            for match in re.findall(r'(Figure\s*\d+(\-\d+)*|Fig\s*\d+(\-\d+)*)', content):
                if ',' in match[0]:
                    [figures_ref.append(int(x.replace('Figure ', '').replace('Fig ', ''))) for x in match[0].split(',')]
                    [curent_ref.append(int(x.replace('Figure ', '').replace('Fig ', ''))) for x in match[0].split(',')]
                    [curent_ref.remove(fig) for fig in curent_ref if fig in checked_fig]
                else:
                    figures_ref.append(int(match[0].replace('Figure ', '').replace('Fig ', '')))
                    curent_ref.append(int(match[0].replace('Figure ', '').replace('Fig ', '')))
                    [curent_ref.remove(fig) for fig in curent_ref if fig in checked_fig]
            curent_ref = list(set(curent_ref))
            refers_count = len(curent_ref)
            if refers_count > 0:
                counter = 0
                for i in range(index + 1, len(document.paragraphs)):
                    if document.paragraphs[i].style.style_id != 'Normal' and document.paragraphs[
                        i].style.style_id != 'FA-Fnumber0' and document.paragraphs[i].style.style_id != 'FA-Break':
                        tab_fig_report['Figure_{0}'.format(curent_ref[counter])] = 'Refer not places correctly'
                    else:
                        if document.paragraphs[i].style.style_id == 'FA-Fnumber0':
                            counter += 1
                    if counter == refers_count:
                        break
                [checked_fig.append(x) for x in curent_ref]
        index += 1
    tables_ref = list(set(tables_ref))
    figures_ref = list(set(figures_ref))
    for key in tables.keys():
        if key not in tables_ref:
            # Table without refer
            tab_fig_report['Table {0}'.format(str(key))] = 'Table without refer'
    for ref in tables_ref:
        if ref not in tables.keys():
            # Table refer has no table
            tab_fig_report['Table {0}'.format(str(ref))] = 'Table refer has no table'
    if len(tab_fig_report) == 0:
        tab_fig_report['Tables'] = 'Ok'
    tab_errors = len(tab_fig_report)
    for key in figures.keys():
        if key not in figures_ref:
            # Figure without refer
            tab_fig_report['Figure {0}'.format(str(key))] = 'Figure without refer'
    for ref in figures_ref:
        if ref not in figures.keys():
            # Figure refer has no figure
            tab_fig_report['Figure {0}'.format(str(ref))] = 'Figure refer has no figure'
    if len(tab_fig_report) == tab_errors:
        tab_fig_report['Figures'] = 'Ok'
    return tab_fig_report


def check_figures_resolutions(article_dir):
    figures = [f for f in os.listdir(article_dir) if f.endswith('.jpg')]
    resolutions = {}
    for figure in figures:
        resolutions[figure] = jpeg_res(os.path.join(article_dir, figure))
    return resolutions


def check_direct_quotations(document):
    quotations_report = {}
    quotations = []
    for paragraph in document.paragraphs:
        quotations.append(re.findall(r'\“(.*?)\”\s*(\(.*?\))', paragraph.text))
        quotations.append(re.findall(r'«(.*?)»\s*(\(.*?\))', paragraph.text))
    quotations = list(filter(None, quotations))
    if len(quotations) == 0:
        quotations_report['Direct quotations'] = 'No quotations'
    for item in quotations:
        citation = item[0][1]
        quotation = item[0][0]
        quotation_count = len(quotation.split())
        found = re.search(r'\(\w*.*?p\.\s+\d+\)|\(\w*.*?pp\.\s+\d+\-\d+\)|\(\w*.*?para\.\s+\d+\)', citation)
        if not found:
            quotations_report[citation] = 'Error, Word count = {0}'.format(quotation_count)
        else:
            quotations_report[citation] = 'Ok, Word count = {0}'.format(quotation_count)
    return quotations_report


def check_references(document):
    ref_report = {}
    references = []
    refers = []
    # All references should be listed alphabetically, with the majority of  current works.
    for para in document.paragraphs:
        if para.style.style_id == 'FA-Paragraphtext':
            for match in re.findall(
                    r'(\(.*?[17|18|19|20]\d{2}\)|(\s*(\w*|\w*\,* et al\.*\,*)\s*(<[^>]*>)*)\s*\(((<[^>]*>)*\s*(17|18|19|20)\d{2}([a-z]{1})*\s*(<[^>]*>)*)\))',
                    para.text):
                content = match[0]
                if content.count('(') > 1:
                    value = content.rfind('(')
                    new_value = content[value:len(content)]
                    refers.append(new_value)
                else:
                    refers.append(content)
        if para.style.style_id == 'FA-RefText':
            references.append(para.text)
    if references != sorted(references):
        ref_report['Order'] = 'Ok'
    else:
        ref_report['Order'] = 'Not Ordered'
    # All references should be cited both in text and in the reference list.
    index = 0
    for para in document.paragraphs:
        if para.style.style_id == 'FA-RefText':
            index += 1
            flag = False
            ref_text = para.text
            for ref in refers:
                parts = ref.split()
                if '&' in parts:
                    parts.remove('&')
                if 'and' in parts:
                    parts.remove('and')
                for word in parts:
                    word = word.replace('(', '').replace(')', '').replace('&', '')
                    if word not in ref_text:
                        flag = False
                        break
                    else:
                        flag = True
                if flag:
                    break
            if not flag:
                ref_report['ref{0}'.format(str(index))] = 'Do not have refer'

    return ref_report


def jpeg_res(filename):
    image_file = Image.open(filename)
    if image_file.info.get('dpi'):
        x_dpi, y_dpi = image_file.info['dpi']
        if x_dpi != y_dpi:
            result = 'Inconsistent DPI image data'
        else:
            result = '{0} dpi'.format(str(x_dpi))
    else:
        result = 'No DPI data. Invalid Image header'
    return result

